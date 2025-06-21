from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
import os
from datetime import datetime
from dotenv import load_dotenv
import openai
from sklearn.metrics.pairwise import cosine_similarity
import requests
import re
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
import io
import tempfile

# Document processing imports
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes

# Load environment variables
load_dotenv()

app = FastAPI(title="Talk to a Folder API")

# OpenAI Configuration
MODEL_NAME = os.getenv("MODEL_NAME", "gpt-4o-mini")
MAX_TOKENS = int(os.getenv("MAX_TOKENS", "1500"))
TEMPERATURE = float(os.getenv("TEMPERATURE", "0.7"))

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

# CORS middleware for frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Vite default port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temporary in-memory storage for MVP
sessions = {}
folder_data = {}
document_store = {}  # Store document chunks and embeddings
embeddings_cache = {}  # Cache for document embeddings
conversation_history = {}  # Store conversation history for each job_id

class AuthRequest(BaseModel):
    access_token: str
    id_token: Optional[str] = None

class IndexRequest(BaseModel):
    access_token: str
    folder_url: str

class ChatRequest(BaseModel):
    access_token: str
    message: str
    job_id: str

class AuthResponse(BaseModel):
    session_id: str

class IndexResponse(BaseModel):
    job_id: str
    status: str
    files_count: int = 0
    folder_name: str = ""

class ChatResponse(BaseModel):
    answer: str
    citations: List[Dict] = []

# Helper functions for AI integration
async def get_embedding(text: str) -> List[float]:
    """Get OpenAI embedding for text"""
    try:
        import openai
        response = openai.Embedding.create(
            model="text-embedding-3-large",
            input=text
        )
        return response['data'][0]['embedding']
    except Exception as e:
        print(f"Error getting embedding: {e}")
        return []

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks"""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
    
    return chunks

async def find_relevant_chunks(query: str, job_id: str, top_k: int = 3) -> List[Dict]:
    """Find most relevant document chunks for a query"""
    if job_id not in document_store:
        return []
    
    # Get query embedding
    query_embedding = await get_embedding(query)
    if not query_embedding:
        return []
    
    # Calculate similarities
    chunks_data = document_store[job_id]
    similarities = []
    
    for chunk_data in chunks_data:
        if chunk_data.get("embedding"):
            similarity = cosine_similarity(
                [query_embedding], 
                [chunk_data["embedding"]]
            )[0][0]
            similarities.append((similarity, chunk_data))
    
    # Sort by similarity and return top_k
    similarities.sort(key=lambda x: x[0], reverse=True)
    return [data for _, data in similarities[:top_k]]

async def generate_answer(query: str, context_chunks: List[Dict], conversation_history: List[Dict] = None) -> str:
    """Generate answer using OpenAI with context and conversation history"""
    try:
        # Prepare context from documents
        context = "\n\n".join([
            f"From {chunk['file_name']}: {chunk['text']}" 
            for chunk in context_chunks
        ])
        
        # Build messages for conversation
        messages = [
            {"role": "system", "content": f"""You are a helpful AI assistant that answers questions about documents in a Google Drive folder. 

Available document context:
{context}

Instructions:
- Use the document context to answer questions accurately
- Reference specific documents when relevant
- If the context doesn't contain enough information, say so honestly
- Maintain conversation continuity by referring to previous messages when appropriate
- Be conversational and helpful"""}
        ]
        
        # Add conversation history (last 10 messages, excluding current user message)
        if conversation_history:
            # Get last 10 messages but exclude the current user message that was just added
            history_to_include = conversation_history[:-1][-10:] if len(conversation_history) > 1 else []
            print(f"🧠 Including {len(history_to_include)} previous messages in context")
            for msg in history_to_include:
                messages.append({
                    "role": msg["role"],
                    "content": msg["content"]
                })
        
        # Add current question
        messages.append({"role": "user", "content": query})

        response = openai.ChatCompletion.create(
            model=MODEL_NAME,
            messages=messages,
            max_tokens=MAX_TOKENS,
            temperature=TEMPERATURE
        )
        
        return response['choices'][0]['message']['content']
        
    except Exception as e:
        print(f"Error generating answer: {e}")
        return f"I apologize, but I encountered an error while processing your question: {str(e)}"

# Google API helper functions
def validate_google_token(access_token: str) -> Dict:
    """Validate Google access token and get user info"""
    try:
        print(f"🔍 Validating Google token (length: {len(access_token) if access_token else 0})")
        print(f"🔍 Token starts with: {access_token[:20]}..." if access_token else "🔍 No token provided")
        
        response = requests.get(
            f'https://www.googleapis.com/oauth2/v1/userinfo?access_token={access_token}'
        )
        
        print(f"🔍 Google API response status: {response.status_code}")
        print(f"🔍 Google API response headers: {dict(response.headers)}")
        
        if response.status_code == 200:
            user_info = response.json()
            print(f"✅ Token validation successful for user: {user_info.get('email', 'unknown')}")
            return user_info
        else:
            error_text = response.text
            print(f"❌ Token validation failed. Status: {response.status_code}, Response: {error_text}")
            raise HTTPException(status_code=401, detail=f"Invalid access token: {error_text}")
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Token validation exception: {str(e)}")
        raise HTTPException(status_code=401, detail=f"Token validation failed: {str(e)}")

def extract_folder_id(folder_url: str) -> str:
    """Extract folder ID from Google Drive URL"""
    # Match various Google Drive folder URL formats
    patterns = [
        r'folders/([a-zA-Z0-9-_]+)',
        r'id=([a-zA-Z0-9-_]+)',
        r'/([a-zA-Z0-9-_]{25,})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, folder_url)
        if match:
            return match.group(1)
    
    raise HTTPException(status_code=400, detail="Invalid Google Drive folder URL")

def get_drive_service(access_token: str):
    """Create Google Drive API service"""
    credentials = Credentials(token=access_token)
    return build('drive', 'v3', credentials=credentials)

# Document processing helper functions
def extract_text_from_pdf(pdf_content: bytes) -> str:
    """Extract text from PDF with fallback to OCR for non-machine readable PDFs"""
    try:
        # First, try standard text extraction
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        text = ""
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            text += page_text + "\n"
        
        # If we got substantial text, return it
        if len(text.strip()) > 100:  # Arbitrary threshold
            print(f"✅ PDF text extraction successful ({len(text)} characters)")
            return text.strip()
        
        # If little or no text found, try OCR
        print("⚠️ PDF appears to be image-based, attempting OCR...")
        return extract_text_from_pdf_ocr(pdf_content)
        
    except Exception as e:
        print(f"❌ PDF text extraction failed, trying OCR: {e}")
        return extract_text_from_pdf_ocr(pdf_content)

def extract_text_from_pdf_ocr(pdf_content: bytes) -> str:
    """Extract text from PDF using OCR"""
    try:
        # Convert PDF pages to images
        images = convert_from_bytes(pdf_content)
        
        text = ""
        for i, image in enumerate(images):
            print(f"🔍 OCR processing page {i + 1}/{len(images)}")
            
            # Use Tesseract OCR to extract text
            page_text = pytesseract.image_to_string(image, lang='eng')
            text += f"\n--- Page {i + 1} ---\n{page_text}\n"
        
        print(f"✅ OCR extraction completed ({len(text)} characters)")
        return text.strip()
        
    except Exception as e:
        print(f"❌ OCR extraction failed: {e}")
        return f"[OCR extraction failed: {str(e)}. Please ensure Tesseract is installed.]"

def extract_text_from_docx(docx_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(docx_content))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        print(f"✅ DOCX text extraction successful ({len(text)} characters)")
        return text.strip()
        
    except Exception as e:
        print(f"❌ DOCX extraction failed: {e}")
        return f"[DOCX extraction failed: {str(e)}]"

def extract_text_from_html(html_content: bytes) -> str:
    """Extract text from HTML file"""
    try:
        # Decode bytes to string
        html_text = html_content.decode('utf-8', errors='ignore')
        
        # Parse HTML and extract text
        soup = BeautifulSoup(html_text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        print(f"✅ HTML text extraction successful ({len(text)} characters)")
        return text
        
    except Exception as e:
        print(f"❌ HTML extraction failed: {e}")
        return f"[HTML extraction failed: {str(e)}]"

async def fetch_folder_files(access_token: str, folder_id: str) -> List[Dict]:
    """Fetch files from Google Drive folder"""
    try:
        print(f"🔍 Fetching files from folder ID: {folder_id}")
        service = get_drive_service(access_token)
        
        # Get folder name
        print(f"🔍 Getting folder info...")
        folder = service.files().get(fileId=folder_id, fields='name').execute()
        folder_name = folder.get('name', 'Unknown Folder')
        print(f"✅ Folder name: {folder_name}")
        
        # Get files in folder
        print(f"🔍 Listing files in folder...")
        query = f"'{folder_id}' in parents and trashed=false"
        print(f"🔍 Query: {query}")
        
        results = service.files().list(
            q=query,
            fields="files(id,name,mimeType,size,modifiedTime)",
            pageSize=100
        ).execute()
        
        files = results.get('files', [])
        print(f"🔍 Found {len(files)} total files")
        
        # Log all files found
        for i, file in enumerate(files):
            print(f"📄 File {i+1}: {file['name']} (type: {file.get('mimeType', 'unknown')})")
        
        # Filter for supported file types
        supported_types = [
            # Text files
            'text/plain',
            'text/csv',
            'text/html',
            'application/xhtml+xml',
            'application/rtf',
            
            # PDF files (with OCR support)
            'application/pdf',
            
            # Microsoft Office documents
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # DOCX
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',         # XLSX
            'application/vnd.openxmlformats-officedocument.presentationml.presentation', # PPTX
            'application/msword',           # Legacy DOC
            'application/vnd.ms-excel',     # Legacy XLS
            'application/vnd.ms-powerpoint', # Legacy PPT
            
            # Google Workspace documents
            'application/vnd.google-apps.document',     # Google Docs
            'application/vnd.google-apps.spreadsheet',  # Google Sheets
            'application/vnd.google-apps.presentation'  # Google Slides
        ]
        
        print(f"🔍 Supported file types: {supported_types}")
        
        filtered_files = []
        for file in files:
            file_type = file.get('mimeType', '')
            if file_type in supported_types:
                print(f"✅ Including file: {file['name']} (type: {file_type})")
                filtered_files.append({
                    'id': file['id'],
                    'name': file['name'],
                    'mimeType': file['mimeType'],
                    'size': file.get('size', 0),
                    'modifiedTime': file.get('modifiedTime')
                })
            else:
                print(f"❌ Skipping file: {file['name']} (unsupported type: {file_type})")
        
        print(f"🔍 Filtered to {len(filtered_files)} supported files")
        
        return {
            'folder_name': folder_name,
            'files': filtered_files
        }
        
    except Exception as e:
        print(f"❌ Error fetching folder files: {e}")
        raise HTTPException(status_code=400, detail=f"Could not access folder: {str(e)}")

async def download_file_content(access_token: str, file_id: str, mime_type: str) -> str:
    """Download and extract text content from a file"""
    try:
        service = get_drive_service(access_token)
        
        print(f"📁 Processing file with MIME type: {mime_type}")
        
        if mime_type == 'application/vnd.google-apps.document':
            # Export Google Doc as plain text
            content = service.files().export(
                fileId=file_id, 
                mimeType='text/plain'
            ).execute()
            text = content.decode('utf-8')
            print(f"✅ Google Doc processed ({len(text)} characters)")
            return text
        
        elif mime_type == 'text/plain':
            # Download plain text file
            content = service.files().get_media(fileId=file_id).execute()
            text = content.decode('utf-8', errors='ignore')
            print(f"✅ TXT file processed ({len(text)} characters)")
            return text
        
        elif mime_type == 'application/pdf':
            # Download PDF file and extract text
            content = service.files().get_media(fileId=file_id).execute()
            return extract_text_from_pdf(content)
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Download DOCX file and extract text
            content = service.files().get_media(fileId=file_id).execute()
            return extract_text_from_docx(content)
        
        elif mime_type in ['text/html', 'application/xhtml+xml']:
            # Download HTML file and extract text
            content = service.files().get_media(fileId=file_id).execute()
            return extract_text_from_html(content)
        
        elif mime_type == 'text/csv':
            # Download CSV file
            content = service.files().get_media(fileId=file_id).execute()
            text = content.decode('utf-8', errors='ignore')
            print(f"✅ CSV file processed ({len(text)} characters)")
            return text
        
        elif mime_type == 'application/rtf':
            # Download RTF file (basic text extraction)
            content = service.files().get_media(fileId=file_id).execute()
            text = content.decode('utf-8', errors='ignore')
            print(f"✅ RTF file processed ({len(text)} characters)")
            return text
        
        elif mime_type == 'application/vnd.google-apps.spreadsheet':
            # Export Google Sheets as CSV
            content = service.files().export(
                fileId=file_id, 
                mimeType='text/csv'
            ).execute()
            text = content.decode('utf-8')
            print(f"✅ Google Sheets processed ({len(text)} characters)")
            return text
        
        elif mime_type == 'application/vnd.google-apps.presentation':
            # Export Google Slides as plain text
            content = service.files().export(
                fileId=file_id, 
                mimeType='text/plain'
            ).execute()
            text = content.decode('utf-8')
            print(f"✅ Google Slides processed ({len(text)} characters)")
            return text
        
        else:
            print(f"⚠️ Unsupported file type: {mime_type}")
            return f"[File type {mime_type} not yet supported for text extraction]"
            
    except Exception as e:
        print(f"❌ Error downloading file {file_id}: {e}")
        return f"[Error reading file: {str(e)}]"

@app.get("/")
async def root():
    return {"message": "Talk to a Folder API is running"}

@app.post("/auth/google", response_model=AuthResponse)
async def auth_google(request: AuthRequest):
    print(f"🔐 Received Google auth request")
    print(f"🔐 Access token provided: {bool(request.access_token)}")
    print(f"🔐 ID token provided: {bool(request.id_token)}")
    
    try:
        # Validate the Google access token
        print(f"🔐 Starting token validation...")
        user_info = validate_google_token(request.access_token)
        print(f"🔐 Token validation completed successfully")
        
        # Create session with real user data
        session_id = f"session_{user_info['id']}_{int(datetime.now().timestamp())}"
        sessions[session_id] = {
            "created_at": datetime.now().isoformat(),
            "access_token": request.access_token,
            "user_info": user_info,
            "authenticated": True
        }
        
        print(f"✅ Session created successfully: {session_id}")
        return AuthResponse(session_id=session_id)
        
    except Exception as e:
        print(f"❌ Auth endpoint error: {str(e)}")
        raise

@app.post("/index", response_model=IndexResponse)
async def index_folder(request: IndexRequest):
    # Validate access token
    _ = validate_google_token(request.access_token)
    
    job_id = f"job_{len(folder_data) + 1}"
    
    try:
        # Extract folder ID from URL
        folder_id = extract_folder_id(request.folder_url)
        
        # Fetch files from Google Drive
        folder_info = await fetch_folder_files(request.access_token, folder_id)
        folder_name = folder_info['folder_name']
        files = folder_info['files']
        
        if not files:
            raise HTTPException(status_code=400, detail="No supported files found in folder")
        
        # Process files and create embeddings
        document_chunks = []
        
        for file in files[:5]:  # Limit to first 5 files for now
            try:
                print(f"Processing file: {file['name']}")
                content = await download_file_content(
                    request.access_token, 
                    file['id'], 
                    file['mimeType']
                )
                
                if content and not content.startswith('['):  # Skip error messages
                    chunks = chunk_text(content)
                    
                    for i, chunk in enumerate(chunks):
                        # Get embedding for each chunk
                        embedding = await get_embedding(chunk)
                        
                        chunk_data = {
                            "file_name": file['name'],
                            "file_id": file['id'],
                            "chunk_id": f"{file['id']}_chunk_{i}",
                            "text": chunk,
                            "embedding": embedding,
                            "mime_type": file['mimeType']
                        }
                        document_chunks.append(chunk_data)
                        
            except Exception as e:
                print(f"Error processing file {file['name']}: {e}")
                continue
        
        if not document_chunks:
            raise HTTPException(status_code=400, detail="Could not extract text from any files")
        
        # Store processed documents
        document_store[job_id] = document_chunks
        
        folder_data[job_id] = {
            "folder_url": request.folder_url,
            "folder_id": folder_id,
            "folder_name": folder_name,
            "status": "completed",
            "files": files,
            "chunks_count": len(document_chunks),
            "access_token": request.access_token
        }
        
        return IndexResponse(
            job_id=job_id, 
            status="completed",
            files_count=len(files),
            folder_name=folder_name
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error processing folder: {e}")
        # Fallback to error status
        folder_data[job_id] = {
            "folder_url": request.folder_url,
            "status": "failed",
            "error": str(e),
            "files": []
        }
        raise HTTPException(status_code=500, detail=f"Failed to process folder: {str(e)}")


@app.get("/index/{job_id}")
async def get_index_status(job_id: str):
    if job_id not in folder_data:
        raise HTTPException(status_code=404, detail="Job not found")
    
    return {
        "job_id": job_id,
        "status": folder_data[job_id]["status"],
        "files_count": len(folder_data[job_id]["files"])
    }

@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    # Validate access token
    _ = validate_google_token(request.access_token)
    
    # Check if the job exists and is completed
    if request.job_id not in folder_data:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job_data = folder_data[request.job_id]
    if job_data.get("status") != "completed":
        return ChatResponse(
            answer="The folder is still being processed. Please wait a moment and try again.",
            citations=[]
        )
    
    try:
        # Initialize conversation history for this job_id if it doesn't exist
        if request.job_id not in conversation_history:
            conversation_history[request.job_id] = []
        
        # Add user message to conversation history
        conversation_history[request.job_id].append({
            "role": "user",
            "content": request.message
        })
        
        # Use RAG pipeline for intelligent responses
        relevant_chunks = await find_relevant_chunks(request.message, request.job_id)
        
        if not relevant_chunks:
            # Fallback if no relevant chunks found
            answer = f"I couldn't find relevant information in the indexed files from '{job_data.get('folder_name', 'your folder')}' to answer that question. Try asking about the content of the documents in the folder."
        else:
            # Generate AI response with context and conversation history
            answer = await generate_answer(
                request.message, 
                relevant_chunks, 
                conversation_history[request.job_id]
            )
        
        # Add AI response to conversation history
        conversation_history[request.job_id].append({
            "role": "assistant",
            "content": answer
        })
        
        # Keep only last 20 messages (10 exchanges) to prevent memory issues
        if len(conversation_history[request.job_id]) > 20:
            conversation_history[request.job_id] = conversation_history[request.job_id][-20:]
        
        # Create citations from relevant chunks
        citations = []
        seen_files = set()
        for chunk in relevant_chunks:
            file_name = chunk["file_name"]
            if file_name not in seen_files:
                citations.append({
                    "file_name": file_name,
                    "file_id": chunk["file_id"],
                    "chunk_id": chunk["chunk_id"]
                })
                seen_files.add(file_name)
        
        print(f"💬 Conversation history for {request.job_id}: {len(conversation_history[request.job_id])} messages")
        
        return ChatResponse(answer=answer, citations=citations)
        
    except Exception as e:
        print(f"Error in chat endpoint: {e}")
        return ChatResponse(
            answer=f"I apologize, but I encountered an error while processing your question. Please make sure your OpenAI API key is properly configured. Error: {str(e)}",
            citations=[]
        )

@app.delete("/chat/{job_id}/history")
async def clear_conversation_history(job_id: str):
    """Clear conversation history for a specific job_id"""
    if job_id in conversation_history:
        del conversation_history[job_id]
        return {"message": f"Conversation history cleared for job {job_id}"}
    return {"message": f"No conversation history found for job {job_id}"}

@app.get("/chat/{job_id}/history")
async def get_conversation_history(job_id: str):
    """Get conversation history for a specific job_id (for debugging)"""
    if job_id in conversation_history:
        return {
            "job_id": job_id,
            "message_count": len(conversation_history[job_id]),
            "messages": conversation_history[job_id]
        }
    return {"job_id": job_id, "message_count": 0, "messages": []}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)